#!/usr/bin/python
# -*- coding:utf-8 -*-
import csv
import docx
import datetime
from setdate import upcoming_weekday # 내가 만든 초간단 라이브러리



# 시간 설정
datenow = datetime.datetime.now()
upcoming_friday = upcoming_weekday(datenow, 4).date() #금요일 설정

# 워드파일 설정
document = docx.Document()

# 엑셀파일 열고 워드파일에 작성
f = open(upcoming_friday.strftime('%Y%m%d')+'_Nat_Neuro.csv', 'r', encoding='utf-8')
rows = csv.reader(f)
for row in rows:
    for s in row:
        # print(s)
        p = document.add_paragraph(s)
    p = document.add_paragraph(' ')
    # for row in rows:
        # print(' '.join(row))
f.close()

document.save(upcoming_friday.strftime('%Y%m%d')+'_Nat_Neuro.docx')
