# Import Libraries
import scrapy
import datetime
from docx import Document

# Time module
now = datetime.datetime.now()
nowDate = now.strftime('%Y%m%d')

# Make word file
document = Document()
document.add_heading(nowDate+' Nature Neuroscience', 0)


# Crawling Nature neuroscience reserch articles
class QuotesSpider(scrapy.Spider):
    name = "quotes"
    start_urls = [
        'https://www.nature.com/neuro/research',
    ]

    def parse(self, response):
        articles = response.xpath('//article')
        for article in articles:
            articletype = article.xpath('.//*[@data-test="article.type"]/text()').extract_first()

            datetime = article.xpath('.//*[@itemprop="datePublished"]/@datetime').extract_first()

            headline = article.xpath('.//*[@itemprop="name headline"]//text()').extract()
            headline = " ".join(headline)
            headline = headline.replace('\n','')
            headline = headline.replace('\t','')
            headline = headline.replace('  ','')

            description = article.xpath('.//*[@itemprop="description"]//p/text()').extract()
            description = " ".join(description)

            if articletype.find('Correction') == -1:
                yield {
                "articletype" : articletype,
                "datetime" : datetime,
                "headline" : headline,
                "description" : description,
                }
                print('\n')
                print(articletype)
                print(datetime)
                print(headline)
                print(description)
                print('\n')

                # document.add_heading(nowDate+' Nature Neuroscience', 0)
                # document.add_paragraph(
                # '\n'+articletype+'\t'+datetime+'\n'+headline+'\n'+description+'\n'
                # )

document.save(nowDate+'_NatNeuro.docx')
